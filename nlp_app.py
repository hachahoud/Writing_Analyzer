import streamlit as st
import spacy
from docx import Document
from datetime import datetime
from pathlib import Path
import re
from spellchecker import SpellChecker
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
import zipfile

# Page configuration
st.set_page_config(
    page_title="Student Writing Development Analyzer",
    page_icon="📝",
    layout="wide"
)

# Load models (cached for performance)
@st.cache_resource
def load_models():
    nlp = spacy.load("en_core_web_sm")
    spell = SpellChecker()
    return nlp, spell

nlp, spell = load_models()

def filter_text(text):
    """Filters and cleans the input text by removing misspelled words"""
    text = text.lower()
    
    contractions = {
        "n't": " not", "'ll": " will", "'ve": " have",
        "'re": " are", "'d": " would", "'m": " am"
    }
    
    for contraction, expansion in contractions.items():
        text = text.replace(contraction, expansion)
    
    text = re.sub(r'[^a-z\s.,!?;:()\']', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    
    words = text.split()
    valid_words = []
    
    for word in words:
        if len(word) == 1 and word not in ['a', 'i']:
            continue
        if word in ['mr', 'mrs', 'ms', 'dr', 'prof', 'etc', 'vs', 'eg', 'ie']:
            valid_words.append(word)
            continue
        if spell.correction(word) == word:
            valid_words.append(word)
    
    filtered_text = ' '.join(valid_words)
    filtered_text = re.sub(r'\s+', ' ', filtered_text).strip()
    
    return filtered_text

def calculate_ttr_spacy(text, segment_length=30):
    """Measures the Type-Token Ratio (TTR), Moving Average Type-Token Ratio (MATTR), and Lexical Density"""
    filtered_text = filter_text(text)
    doc = nlp(filtered_text.lower())
    
    words = [token.text for token in doc if token.is_alpha]
    unique_words = set(words)
    ttr = len(unique_words) / len(words) if words else 0
    
    if len(words) < segment_length:
        mattr = ttr
    else:
        ttr_values = []
        for i in range(len(words) - segment_length + 1):
            window = words[i:i + segment_length]
            window_ttr = len(set(window)) / len(window)
            ttr_values.append(window_ttr)
        mattr = sum(ttr_values) / len(ttr_values)

    lexical_words = [token.text for token in doc if token.is_alpha and token.pos_ in ['NOUN', 'VERB', 'ADJ', 'ADV']]
    lexical_density = len(lexical_words) / len(words) if words else 0

    return ttr, mattr, len(words), len(unique_words), lexical_density

def calculate_dcr(text):
    """Calculates the ratio of dependent clauses to independent clauses in the text"""
    filtered_text = filter_text(text)
    doc = nlp(filtered_text)
    
    dependent_clauses = 0
    independent_clauses = 0

    for token in doc:
        if token.dep_ in ["ccomp", "advcl", "acl", "relcl", "xcomp", "pcomp"]:
            dependent_clauses += 1
    
    for sent in doc.sents:
        root = [token for token in sent if token.dep_ == "ROOT"]
        if root:
            root = root[0]
            
            def get_coordinated_clauses(token):
                clauses = []
                for child in token.children:
                    if child.dep_ == "conj":
                        clauses.append(child)
                        clauses.extend(get_coordinated_clauses(child))
                return clauses
            
            coordinated_clauses = get_coordinated_clauses(root)
            independent_clauses += 1
            independent_clauses += len(coordinated_clauses)
    
    total_clauses = independent_clauses + dependent_clauses
    dcr = dependent_clauses / total_clauses if total_clauses > 0 else 0
    
    return dcr, dependent_clauses, total_clauses

def analyze_text(text, date_label):
    """Analyzes a single text sample"""
    ttr, mattr, total_words, unique_words, lexical_density = calculate_ttr_spacy(text)
    dcr, dependent_clauses, total_clauses = calculate_dcr(text)
    
    return {
        'date': date_label,
        'ttr': ttr,
        'mattr': mattr,
        'total_words': total_words,
        'unique_words': unique_words,
        'lexical_density': lexical_density,
        'dcr': dcr,
        'dependent_clauses': dependent_clauses,
        'total_clauses': total_clauses
    }

def create_word_report(student_name, results_df):
    """Creates a Word document with analysis results"""
    doc = Document()
    doc.add_heading(f'Writing Development Analysis - {student_name}', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    table = doc.add_table(rows=len(results_df) + 1, cols=10)
    table.style = 'Table Grid'
    
    headers = ['Date', 'Total Words', 'Unique Words', 'TTR', 'MATTR', 'Lexical Density', 
              'DCR', 'Dependent Clauses', 'Total Clauses', 'File Name']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    for i, row in enumerate(results_df.itertuples(), start=1):
        table.cell(i, 0).text = str(row.date)
        table.cell(i, 1).text = str(row.total_words)
        table.cell(i, 2).text = str(row.unique_words)
        table.cell(i, 3).text = f"{row.ttr:.2f} ({row.ttr*100:.2f}%)"
        table.cell(i, 4).text = f"{row.mattr:.2f} ({row.mattr*100:.2f}%)"
        table.cell(i, 5).text = f"{row.lexical_density:.2f} ({row.lexical_density*100:.2f}%)"
        table.cell(i, 6).text = f"{row.dcr:.2f} ({row.dcr*100:.2f}%)"
        table.cell(i, 7).text = str(row.dependent_clauses)
        table.cell(i, 8).text = str(row.total_clauses)
        table.cell(i, 9).text = str(row.file_name)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.title("📝 Student Writing Development Analyzer")
st.markdown("Analyze writing samples to track linguistic development over time")

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    analysis_mode = st.radio(
        "Choose Analysis Mode:",
        ["Single Text Analysis", "Multiple Files (Time Series)"]
    )
    
    st.markdown("---")
    st.markdown("### 📊 Metrics Explained")
    with st.expander("What is TTR?"):
        st.write("**Type-Token Ratio** measures vocabulary diversity. Higher values indicate more varied vocabulary.")
    with st.expander("What is MATTR?"):
        st.write("**Moving Average Type-Token Ratio** is a more stable measure of vocabulary diversity using sliding windows.")
    with st.expander("What is DCR?"):
        st.write("**Dependent Clause Ratio** measures syntactic complexity. Higher values indicate more complex sentence structures.")
    with st.expander("What is Lexical Density?"):
        st.write("**Lexical Density** measures the proportion of content words (nouns, verbs, adjectives, adverbs) in the text.")

# Main content
if analysis_mode == "Single Text Analysis":
    st.header("✍️ Single Text Analysis")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        text_input = st.text_area(
            "Enter text to analyze:",
            height=300,
            placeholder="Paste the student's writing here..."
        )
    
    with col2:
        date_label = st.text_input("Date/Label (optional):", value=datetime.now().strftime("%Y-%m-%d"))
        student_name = st.text_input("Student Name:", value="Student")
    
    if st.button("🔍 Analyze Text", type="primary"):
        if text_input.strip():
            with st.spinner("Analyzing text..."):
                result = analyze_text(text_input, date_label)
                
                st.success("✅ Analysis Complete!")
                
                # Display metrics in columns
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total Words", result['total_words'])
                    st.metric("Unique Words", result['unique_words'])
                with col2:
                    st.metric("TTR", f"{result['ttr']:.2%}")
                    st.metric("MATTR", f"{result['mattr']:.2%}")
                with col3:
                    st.metric("Lexical Density", f"{result['lexical_density']:.2%}")
                    st.metric("DCR", f"{result['dcr']:.2%}")
                with col4:
                    st.metric("Dependent Clauses", result['dependent_clauses'])
                    st.metric("Total Clauses", result['total_clauses'])
        else:
            st.warning("⚠️ Please enter some text to analyze.")

else:  # Multiple Files Analysis
    st.header("📁 Multiple Files Analysis")
    
    student_name = st.text_input("Student Name:", value="Student")
    
    uploaded_files = st.file_uploader(
        "Upload text files (.txt)",
        type=['txt'],
        accept_multiple_files=True,
        help="Upload multiple text files. Filename should include the date (e.g., 2024-01-15.txt, week1.txt, etc)"
    )
    
    if uploaded_files:
        st.info(f"📄 {len(uploaded_files)} file(s) uploaded")
        
        if st.button("🔍 Analyze All Files", type="primary"):
            results = []
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, file in enumerate(uploaded_files):
                status_text.text(f"Analyzing {file.name}...")
                text = file.read().decode('utf-8')
                date_label = Path(file.name).stem
                
                result = analyze_text(text, date_label)
                result['file_name'] = file.name
                results.append(result)
                
                progress_bar.progress((idx + 1) / len(uploaded_files))
            
            status_text.text("✅ Analysis complete!")
            
            # Convert to DataFrame
            df = pd.DataFrame(results)
            df = df.sort_values('date')
            
            # Store in session state
            st.session_state['results_df'] = df
            st.session_state['student_name'] = student_name
            
            # Display results
            st.success(f"✅ Analyzed {len(results)} files!")
            
            # Summary Statistics
            st.subheader("📊 Summary Statistics")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Avg Total Words", f"{df['total_words'].mean():.0f}")
                st.metric("Avg Unique Words", f"{df['unique_words'].mean():.0f}")
            with col2:
                st.metric("Avg TTR", f"{df['ttr'].mean():.2%}")
                st.metric("Avg MATTR", f"{df['mattr'].mean():.2%}")
            with col3:
                st.metric("Avg Lexical Density", f"{df['lexical_density'].mean():.2%}")
                st.metric("Avg DCR", f"{df['dcr'].mean():.2%}")
            with col4:
                st.metric("Total Samples", len(df))
                st.metric("Date Range", f"{df['date'].min()} to {df['date'].max()}")
            
            # Detailed Results Table
            st.subheader("📋 Detailed Results")
            st.dataframe(
                df.style.format({
                    'ttr': '{:.2%}',
                    'mattr': '{:.2%}',
                    'lexical_density': '{:.2%}',
                    'dcr': '{:.2%}'
                }),
                width='stretch'
            )
            
            # Visualizations
            st.subheader("📈 Progress Over Time")
            
            tab1, tab2, tab3, tab4 = st.tabs(["Vocabulary", "Complexity", "Lexical Density", "Word Count"])
            
            with tab1:
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=df['date'], y=df['ttr'], mode='lines+markers', name='TTR'))
                fig.add_trace(go.Scatter(x=df['date'], y=df['mattr'], mode='lines+markers', name='MATTR'))
                fig.update_layout(title="Vocabulary Diversity Over Time", xaxis_title="Date", yaxis_title="Ratio")
                st.plotly_chart(fig, width='stretch')

            with tab2:
                fig = px.line(df, x='date', y='dcr', markers=True, title="Syntactic Complexity (DCR) Over Time")
                fig.update_layout(xaxis_title="Date", yaxis_title="DCR")
                st.plotly_chart(fig, width='stretch')
            
            with tab3:
                fig = px.line(df, x='date', y='lexical_density', markers=True, title="Lexical Density Over Time")
                fig.update_layout(xaxis_title="Date", yaxis_title="Lexical Density")
                st.plotly_chart(fig, width='stretch')
            
            with tab4:
                fig = go.Figure()
                fig.add_trace(go.Bar(x=df['date'], y=df['total_words'], name='Total Words'))
                fig.add_trace(go.Bar(x=df['date'], y=df['unique_words'], name='Unique Words'))
                fig.update_layout(title="Word Count Over Time", xaxis_title="Date", yaxis_title="Count", barmode='group')
                st.plotly_chart(fig, width='stretch')
            
            # Download Section
            st.subheader("💾 Download Reports")
            col1, col2 = st.columns(2)
            
            with col1:
                # CSV Download
                csv = df.to_csv(index=False)
                st.download_button(
                    label="📥 Download CSV",
                    data=csv,
                    file_name=f"{student_name}_analysis.csv",
                    mime="text/csv"
                )
            
            with col2:
                # Word Document Download
                word_buffer = create_word_report(student_name, df)
                st.download_button(
                    label="📥 Download Word Report",
                    data=word_buffer,
                    file_name=f"{student_name}_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Footer
st.markdown("---")
st.markdown("*Created by Hamza Achahoud © 2025*")